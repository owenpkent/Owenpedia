"""Script to convert XLSX files to CSV and DOCX files to Markdown."""
import os
from pathlib import Path
import pandas as pd
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

def xlsx_to_csv(xlsx_path, output_dir):
    """Convert Excel file to CSV(s). Multiple sheets create multiple CSVs."""
    xlsx_path = Path(xlsx_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    try:
        xl = pd.ExcelFile(xlsx_path)
        base_name = xlsx_path.stem
        
        for sheet_name in xl.sheet_names:
            df = pd.read_excel(xl, sheet_name=sheet_name)
            if len(xl.sheet_names) > 1:
                csv_name = f"{base_name}_{sheet_name}.csv"
            else:
                csv_name = f"{base_name}.csv"
            csv_path = output_dir / csv_name
            df.to_csv(csv_path, index=False)
            print(f"  Created: {csv_path}")
    except Exception as e:
        print(f"  ERROR converting {xlsx_path}: {e}")

def docx_to_markdown(docx_path, output_dir):
    """Convert Word document to Markdown."""
    docx_path = Path(docx_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    try:
        doc = Document(docx_path)
        md_lines = []
        
        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('p'):
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if not text:
                            md_lines.append("")
                            continue
                        
                        style = para.style.name if para.style else ""
                        
                        # Handle headings
                        if "Heading 1" in style or "Title" in style:
                            md_lines.append(f"# {text}")
                        elif "Heading 2" in style:
                            md_lines.append(f"## {text}")
                        elif "Heading 3" in style:
                            md_lines.append(f"### {text}")
                        elif "List" in style or text.startswith(("•", "-", "*", "·")):
                            clean_text = text.lstrip("•-*·→ ")
                            md_lines.append(f"- {clean_text}")
                        else:
                            # Check for bold/emphasis
                            md_lines.append(text)
                        break
            
            # Handle tables
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._tbl == element:
                        md_lines.append("")
                        if table.rows:
                            # Header row
                            headers = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
                            md_lines.append("| " + " | ".join(headers) + " |")
                            md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                            
                            # Data rows
                            for row in table.rows[1:]:
                                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                                md_lines.append("| " + " | ".join(cells) + " |")
                        md_lines.append("")
                        break
        
        # Clean up multiple blank lines
        cleaned_lines = []
        prev_blank = False
        for line in md_lines:
            is_blank = line == ""
            if is_blank and prev_blank:
                continue
            cleaned_lines.append(line)
            prev_blank = is_blank
        
        md_content = "\n".join(cleaned_lines)
        md_path = output_dir / f"{docx_path.stem}.md"
        md_path.write_text(md_content, encoding='utf-8')
        print(f"  Created: {md_path}")
        
    except Exception as e:
        print(f"  ERROR converting {docx_path}: {e}")

def main():
    repo_root = Path(__file__).parent.parent
    
    # Find all xlsx and docx files
    xlsx_files = list(repo_root.glob("**/*.xlsx"))
    docx_files = list(repo_root.glob("**/*.docx"))
    
    print(f"Found {len(xlsx_files)} Excel files and {len(docx_files)} Word documents\n")
    
    print("=== Converting Excel files to CSV ===")
    for xlsx_file in xlsx_files:
        print(f"Processing: {xlsx_file.name}")
        # Keep same relative directory structure
        rel_dir = xlsx_file.parent.relative_to(repo_root)
        xlsx_to_csv(xlsx_file, repo_root / rel_dir)
    
    print("\n=== Converting Word documents to Markdown ===")
    for docx_file in docx_files:
        print(f"Processing: {docx_file.name}")
        rel_dir = docx_file.parent.relative_to(repo_root)
        docx_to_markdown(docx_file, repo_root / rel_dir)
    
    print("\n=== Conversion complete ===")

if __name__ == "__main__":
    main()
